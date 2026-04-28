#!/usr/bin/env python3
"""
Convert paper_31_draft.md to Paper_31.docx via python-docx.

Handles the specific markdown subset used in the paper:
  - ATX headings (# through ####)
  - Paragraphs with **bold**, *italic*, `code`
  - Tables (GFM pipe style)
  - Bullet lists (- )
  - Numbered lists (1. 2. 3.)
  - Horizontal rules (---) as section separators
  - Image embed (![](path)) -> inline picture
  - Figure captions handled as italic paragraphs below the image

Intentionally does NOT handle: code blocks (>4 spaces or fenced), math ($...$),
reference-style links, blockquotes, or nested lists.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).parent
# Defaults; override via CLI: `python md_to_docx.py INPUT.md OUTPUT.docx`
MD_PATH    = SCRIPT_DIR / "paper_31_draft.md"
DOCX_PATH  = SCRIPT_DIR / "Paper_31.docx"
import sys
if len(sys.argv) >= 3:
    MD_PATH   = Path(sys.argv[1]) if Path(sys.argv[1]).is_absolute() else SCRIPT_DIR / sys.argv[1]
    DOCX_PATH = Path(sys.argv[2]) if Path(sys.argv[2]).is_absolute() else SCRIPT_DIR / sys.argv[2]


# ---------------------------------------------------------------------------
# Inline formatting parser
# ---------------------------------------------------------------------------
def add_runs_from_inline(paragraph, text):
    """Parse inline markdown (**bold**, *italic*, `code`) and append runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            r = paragraph.add_run(token[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------------------------------------------------------------------------
# Helpers for page setup
# ---------------------------------------------------------------------------
def apply_paragraph_style(p, *, font="Calibri", size=11, bold=False,
                          italic=False, space_before=0, space_after=6,
                          align=None):
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align is not None:
        p.alignment = align
    for r in p.runs:
        r.font.name = font
        r.font.size = Pt(size)
        if bold:   r.bold = True
        if italic: r.italic = True


# ---------------------------------------------------------------------------
# Block parser
# ---------------------------------------------------------------------------
def parse_md(md_text: str):
    """Yield block tuples: ('h1', text), ('h2', text), ('p', text),
    ('hr',), ('table', rows), ('ul', items), ('ol', items),
    ('image', path, alt)."""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i]
        s = ln.rstrip()

        if s.strip() == "":
            i += 1; continue

        if s.strip() == "---":
            yield ("hr",); i += 1; continue

        # Fenced code block: ``` ... ```
        m = re.match(r"^\s*```(\w*)\s*$", s)
        if m:
            lang = m.group(1)
            code_lines = []
            i += 1
            while i < len(lines) and not re.match(r"^\s*```\s*$", lines[i]):
                code_lines.append(lines[i])
                i += 1
            # Skip the closing fence
            if i < len(lines):
                i += 1
            yield ("code", code_lines, lang)
            continue

        # ATX headings
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            yield (f"h{len(m.group(1))}", m.group(2).strip())
            i += 1; continue

        # Image — !(alt)[path] or ![alt](path)
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", s)
        if m:
            yield ("image", m.group(2), m.group(1))
            i += 1; continue

        # Table
        if "|" in s and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+:?", lines[i + 1]):
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(lines[i])
                i += 1
            if len(rows) >= 2:
                yield ("table", rows)
                continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", s):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            yield ("ul", items)
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", s):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            yield ("ol", items)
            continue

        # Regular paragraph — collect until blank line
        buf = [s]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#"):
            if re.match(r"^\s*\|", lines[i]):
                break
            buf.append(lines[i])
            i += 1
        yield ("p", " ".join(buf))


# ---------------------------------------------------------------------------
# Renderer
# ---------------------------------------------------------------------------
def parse_table_rows(rows):
    parsed = []
    for r in rows:
        r = r.strip().strip("|")
        cells = [c.strip() for c in re.split(r"\s*\|\s*", r)]
        parsed.append(cells)
    # Drop divider row (row 1)
    header = parsed[0]
    data   = [p for p in parsed[2:] if p]
    return header, data


def render(md_text: str) -> Document:
    doc = Document()

    # Global style: Calibri 11
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Page margins
    for section in doc.sections:
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    for block in parse_md(md_text):
        tag = block[0]

        if tag == "h1":
            p = doc.add_heading(block[1], level=0)
            for r in p.runs:
                r.font.size = Pt(16)
        elif tag == "h2":
            p = doc.add_heading(block[1], level=1)
            for r in p.runs:
                r.font.size = Pt(14)
        elif tag == "h3":
            p = doc.add_heading(block[1], level=2)
            for r in p.runs:
                r.font.size = Pt(12)
        elif tag == "h4":
            p = doc.add_heading(block[1], level=3)
            for r in p.runs:
                r.font.size = Pt(11)

        elif tag == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            # Insert a horizontal line as a bottom border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),  "single")
            bottom.set(qn("w:sz"),   "6")
            bottom.set(qn("w:color"),"888888")
            bottom.set(qn("w:space"),"1")
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif tag == "code":
            # Fenced code block: render each line as a monospace paragraph
            # with a subtle light-gray shading via paragraph XML.
            code_lines = block[1] if len(block) > 1 else []
            for line in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                # Shade the paragraph background light gray
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "F5F5F5")
                pPr.append(shd)
                # Render text in monospace
                run = p.add_run(line if line else "\u00A0")
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            # Add a blank paragraph after the code block for spacing
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

        elif tag == "image":
            path = SCRIPT_DIR / block[1]
            if not path.exists():
                path = Path(block[1])
            try:
                doc.add_picture(str(path), width=Inches(6.3))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Emit caption (alt text) as italic paragraph below the image
                alt = block[2] if len(block) > 2 else ""
                if alt and alt.strip():
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cap.add_run(alt.strip())
                    run.italic = True
                    run.font.size = Pt(9)
            except Exception as e:
                doc.add_paragraph(f"[image not embedded: {path} -- {e}]")

        elif tag == "p":
            p = doc.add_paragraph()
            add_runs_from_inline(p, block[1])
            apply_paragraph_style(p, space_after=6)

        elif tag == "ul":
            for item in block[1]:
                p = doc.add_paragraph(style="List Bullet")
                add_runs_from_inline(p, item)

        elif tag == "ol":
            for item in block[1]:
                p = doc.add_paragraph(style="List Number")
                add_runs_from_inline(p, item)

        elif tag == "table":
            header, data = parse_table_rows(block[1])
            n_cols = len(header)
            table = doc.add_table(rows=1 + len(data), cols=n_cols)
            table.style = "Light Grid Accent 1"
            # Header
            hdr_cells = table.rows[0].cells
            for c, h in zip(hdr_cells, header):
                c.text = ""
                add_runs_from_inline(c.paragraphs[0], h)
                for r in c.paragraphs[0].runs:
                    r.bold = True
                    r.font.size = Pt(10)
            # Body
            for ri, row in enumerate(data, start=1):
                row_cells = table.rows[ri].cells
                for ci in range(n_cols):
                    cell_text = row[ci] if ci < len(row) else ""
                    row_cells[ci].text = ""
                    add_runs_from_inline(row_cells[ci].paragraphs[0], cell_text)
                    for r in row_cells[ci].paragraphs[0].runs:
                        r.font.size = Pt(10)

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    md = MD_PATH.read_text(encoding="utf-8")
    doc = render(md)
    doc.save(str(DOCX_PATH))
    print(f"wrote {DOCX_PATH} ({DOCX_PATH.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
